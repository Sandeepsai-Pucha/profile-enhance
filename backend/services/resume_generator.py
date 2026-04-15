"""
services/resume_generator.py
─────────────────────────────
Generates a resume .docx in ABSYZ format.

Header fields (Name, Title, Company, Email) — all Calibri 12pt bold, tight spacing.
Company is always "ABSYZ Software Consulting Pvt Ltd".
Profile Summary — rendered as bullet points exactly as parsed.
Technical Skills — new skills appended in plain black, no color, no [NEW] tag.
"""

import re
import io
from typing import List, Dict, Tuple

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


COMPANY_NAME = "ABSYZ Software Consulting Pvt Ltd"
HEADER_FONT  = "Calibri"
HEADER_SIZE  = Pt(11)
BODY_FONT    = "Calibri"
BODY_SIZE    = Pt(11)


# ── Skill → category mapping ──────────────────────────────────
_CATEGORY_KEYWORDS: List[Tuple[str, List[str]]] = [
    ("Frontend", [
        "react", "vue", "angular", "next", "nuxt", "svelte",
        "html", "css", "bootstrap", "tailwind", "sass", "scss",
        "redux", "webpack", "vite", "jquery", "gatsby",
        "material ui", "material-ui", "mui", "ant design",
        "responsive", "web development", "ui", "ux",
    ]),
    ("Backend", [
        "node", "express", "nestjs", "fastapi", "django", "flask",
        "spring", "php", "laravel", "yii", "codeigniter", "ruby",
        "rails", "go", "golang", "rust", "graphql", "grpc",
        "typeorm", "sequelize", "prisma", "serverless", "twilio", "sendgrid",
    ]),
    ("Languages", [
        "python", "javascript", "typescript", "java", "c#", ".net",
        "kotlin", "swift", "dart", "flutter", "c++", "scala", "perl",
    ]),
    ("Databases", [
        "mysql", "postgresql", "postgres", "mongodb", "redis",
        "dynamodb", "sqlite", "oracle", "mssql", "sql server",
        "elasticsearch", "mariadb", "sql",
    ]),
    ("Cloud", [
        "aws", "gcp", "azure", "lambda", "ec2", "s3", "cloudwatch",
        "api gateway", "amplify", "cloudfront", "sns", "sqs",
        "heroku", "netlify", "vercel", "firebase",
    ]),
    ("DevOps & Tools", [
        "docker", "kubernetes", "k8s", "terraform", "ansible", "helm",
        "jenkins", "github actions", "circleci", "ci/cd", "nginx",
        "linux", "bash", "git", "github", "gitlab", "jira", "postman",
    ]),
    ("Testing", [
        "jest", "pytest", "mocha", "chai", "cypress", "selenium",
        "playwright", "junit",
    ]),
    ("AI / ML & Analytics", [
        "openai", "langchain", "llm", "machine learning", "tensorflow",
        "pytorch", "scikit", "pandas", "numpy", "power bi",
    ]),
    ("Architecture & Concepts", [
        "microservices", "rest", "restful", "distributed", "event-driven",
        "api design", "oops", "oop", "agile", "scrum",
    ]),
]


def _categorize(skill: str) -> str:
    low = skill.lower()
    for cat, kws in _CATEGORY_KEYWORDS:
        if any(kw in low for kw in kws):
            return cat
    return "Other"


def _group_skills(existing: List[str], new_added: List[str]) -> List[Dict]:
    cat_ex: Dict[str, List[str]] = {}
    for s in existing:
        cat_ex.setdefault(_categorize(s), []).append(s)

    existing_lower = {s.lower() for s in existing}
    cat_new: Dict[str, List[str]] = {}
    for s in new_added:
        if s.lower() not in existing_lower:
            cat_new.setdefault(_categorize(s), []).append(s)

    all_cats = list(cat_ex.keys())
    for c in cat_new:
        if c not in all_cats:
            all_cats.append(c)

    return [
        {"cat": c, "existing": cat_ex.get(c, []), "new": cat_new.get(c, [])}
        for c in all_cats
        if cat_ex.get(c) or cat_new.get(c)
    ]


# ── Helpers ───────────────────────────────────────────────────

def _add_page_border(doc: Document):
    """
    Add a box border to every section — Dark Blue Text 2 (1F3864), 3pt (36 eighths-of-a-point).
    """
    for section in doc.sections:
        sectPr = section._sectPr
        # Remove existing pgBorders if any
        for existing in sectPr.findall(qn("w:pgBorders")):
            sectPr.remove(existing)

        pgBorders = OxmlElement("w:pgBorders")
        pgBorders.set(qn("w:offsetFrom"), "page")

        for side in ("top", "left", "bottom", "right"):
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"),   "single")
            border.set(qn("w:sz"),    "36")        # 36 × 1/8 pt = 4.5 pt  (Word maps 3pt → sz=24; use 24)
            border.set(qn("w:sz"),    "24")        # 3pt = 24 eighths-of-a-point
            border.set(qn("w:space"), "24")
            border.set(qn("w:color"), "1F3864")    # Dark Blue, Text 2
            pgBorders.append(border)

        sectPr.append(pgBorders)


def _add_email_hyperlink(doc: Document, email: str):
    """Add email as a mailto: hyperlink in blue, bold, Calibri 11pt."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0.5)

    # Register the relationship
    part = doc.part
    r_id = part.relate_to(
        f"mailto:{email}",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    # Build hyperlink XML element
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_elem = OxmlElement("w:r")

    rPr = OxmlElement("w:rPr")
    # Bold
    b = OxmlElement("w:b"); rPr.append(b)
    # Blue colour
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0000FF"); rPr.append(color)
    # Font
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), HEADER_FONT)
    rFonts.set(qn("w:hAnsi"), HEADER_FONT)
    rPr.append(rFonts)
    # Size
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "22"); rPr.append(sz)   # 22 half-pts = 11pt

    run_elem.append(rPr)

    t = OxmlElement("w:t")
    t.text = email
    run_elem.append(t)

    hyperlink.append(run_elem)
    p._p.append(hyperlink)
    return p


def _header_line(doc: Document, text: str):
    """Single header line — Calibri 12pt bold, minimal spacing."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0.5)
    run = p.add_run(text)
    run.bold           = True
    run.font.name      = HEADER_FONT
    run.font.size      = HEADER_SIZE
    run.font.color.rgb = RGBColor(0, 0, 0)
    # Also set the theme font via XML so Word respects Calibri
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"),    HEADER_FONT)
    rFonts.set(qn("w:hAnsi"),    HEADER_FONT)
    rFonts.set(qn("w:eastAsia"), HEADER_FONT)
    return p


def _section_heading(doc: Document, title: str):
    """Bold underlined section heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(title)
    run.bold      = True
    run.underline = True
    run.font.name = BODY_FONT
    run.font.size = BODY_SIZE
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def _add_bullet(doc: Document, text: str, indent_level: int = 0):
    """Bullet point paragraph."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Inches(0.25 * (indent_level + 1))
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = BODY_SIZE
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def _add_field_line(doc: Document, label: str, value: str):
    """Bold label: normal value."""
    if not value:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    bold_run = p.add_run(f"{label}: ")
    bold_run.bold      = True
    bold_run.font.name = BODY_FONT
    bold_run.font.size = BODY_SIZE
    val_run = p.add_run(value)
    val_run.font.name = BODY_FONT
    val_run.font.size = BODY_SIZE
    return p


def _add_skill_row(doc: Document, cat: str, existing: List[str], new_added: List[str]):
    """- Bold Category: all skills in plain black (no color, no [NEW])."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(1)

    cat_run = p.add_run(f"{cat}: ")
    cat_run.bold      = True
    cat_run.font.name = BODY_FONT
    cat_run.font.size = BODY_SIZE

    deduped_new = [s for s in new_added if s.lower() not in {e.lower() for e in existing}]
    all_skills  = existing + deduped_new

    skills_run = p.add_run(", ".join(all_skills))
    skills_run.font.name      = BODY_FONT
    skills_run.font.size      = BODY_SIZE
    skills_run.font.color.rgb = RGBColor(0, 0, 0)


# ── Public builder ────────────────────────────────────────────
def build_resume_docx(page_index: dict, new_skills: List[str]) -> bytes:
    identity  = page_index.get("identity",  {})
    skills_d  = page_index.get("skills",    {})
    exp       = page_index.get("experience", {})
    edu       = page_index.get("education", {})
    narrative = page_index.get("narrative", {})

    name            = identity.get("name", "Candidate") or "Candidate"
    email           = identity.get("email") or ""
    current_role    = identity.get("current_role") or ""
    timeline        = exp.get("timeline") or []
    edu_text        = edu.get("summary") or ""
    certs           = edu.get("certifications") or []
    existing_skills = skills_d.get("all") or []

    # Profile summary — prefer bullet points array, fall back to splitting paragraph
    summary_points: List[str] = narrative.get("summary_points") or []
    if not summary_points:
        raw_summary = narrative.get("summary") or ""
        if raw_summary:
            pts = re.split(r"\.\s+", raw_summary.strip())
            summary_points = [p.strip().rstrip(".") for p in pts if len(p.strip()) > 10]
            if not summary_points:
                summary_points = [raw_summary]

    skill_rows = _group_skills(existing_skills, new_skills)

    doc = Document()

    # Narrow margins
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)

    # Page border — Dark Blue Text 2, 3pt
    _add_page_border(doc)

    # ── HEADER — Name | Title | Company | Email ───────────────
    _header_line(doc, name)
    if current_role:
        _header_line(doc, current_role)
    _header_line(doc, COMPANY_NAME)
    if email:
        _add_email_hyperlink(doc, email)

    # ── PROFILE SUMMARY ───────────────────────────────────────
    if summary_points:
        _section_heading(doc, "PROFILE SUMMARY:")
        for pt in summary_points:
            _add_bullet(doc, pt)

    # ── TECHNICAL SKILLS ─────────────────────────────────────
    if skill_rows:
        _section_heading(doc, "TECHNICAL SKILLS:")
        for row in skill_rows:
            _add_skill_row(doc, row["cat"], row["existing"], row["new"])

    # ── KEY PROJECTS ─────────────────────────────────────────
    if timeline:
        _section_heading(doc, "KEY PROJECTS:")
        for i, job in enumerate(timeline, 1):
            ph = doc.add_paragraph()
            ph.paragraph_format.space_before = Pt(4)
            ph.paragraph_format.space_after  = Pt(1)
            pr = ph.add_run(f"Project #{i}")
            pr.bold      = True
            pr.font.name = BODY_FONT
            pr.font.size = BODY_SIZE

            _add_field_line(doc, "Client Industry", job.get("company") or "")
            _add_field_line(doc, "Role",            job.get("title") or "")
            _add_field_line(doc, "Technologies",    job.get("technologies") or "")
            _add_field_line(doc, "Duration",        job.get("duration") or "")
            _add_field_line(doc, "Description",     job.get("description") or "")

            resps = job.get("key_responsibilities") or []
            if resps:
                rp = doc.add_paragraph()
                rp.paragraph_format.space_after = Pt(1)
                rr = rp.add_run("Responsibilities:")
                rr.bold      = True
                rr.font.name = BODY_FONT
                rr.font.size = BODY_SIZE
                for resp in resps:
                    _add_bullet(doc, resp, indent_level=1)

    # ── EDUCATION & CERTIFICATIONS ────────────────────────────
    if edu_text or certs:
        _section_heading(doc, "EDUCATION & CERTIFICATIONS:")
        if edu_text:
            ep = doc.add_paragraph()
            ep.paragraph_format.space_after = Pt(1)
            er = ep.add_run(edu_text)
            er.font.name = BODY_FONT
            er.font.size = BODY_SIZE
        for cert in certs:
            _add_bullet(doc, cert)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
